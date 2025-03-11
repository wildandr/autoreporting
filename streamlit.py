import streamlit as st
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import requests
import io
import base64
import os
import subprocess
import tempfile

# Set page title
st.set_page_config(page_title="Daily Report Generator", layout="wide")
st.title("Daily Report Generator - Wildan Dzaky Ramadhani")

# Fungsi untuk menyalin pemformatan warna sel
def copy_cell_formatting(src_cell, dest_cell):
    tcPr = src_cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is not None:
        new_tcPr = dest_cell._tc.get_or_add_tcPr()
        new_shd = OxmlElement('w:shd')
        new_shd.set(qn('w:fill'), shd.get(qn('w:fill')))
        new_shd.set(qn('w:val'), shd.get(qn('w:val')))
        new_tcPr.append(new_shd)

# Fungsi untuk menambahkan garis hitam pada sel
def add_black_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # Tipe garis: single
        border.set(qn('w:sz'), '4')       # Ukuran garis: 4 (1/8 pt, cukup tebal)
        border.set(qn('w:space'), '0')    # Jarak dari teks: 0
        border.set(qn('w:color'), '000000')  # Warna: hitam (RGB: 000000)
        tcBorders.append(border)
    tcPr.append(tcBorders)

# Fungsi untuk mengubah format tanggal
def format_tanggal(tanggal_str, bulan_dict):
    if isinstance(tanggal_str, str) and tanggal_str.strip():
        try:
            # Check for "NaT" string or empty string
            if tanggal_str == "NaT" or tanggal_str == "":
                return ""
                
            dt = datetime.strptime(tanggal_str.split()[0], "%Y-%m-%d")
            hari = dt.strftime("%d")
            bulan = bulan_dict[dt.strftime("%B")]
            tahun = dt.strftime("%Y")
            return f"{hari} {bulan} {tahun}"
        except ValueError:
            # Return empty string for invalid dates
            return ""
    # Return empty string for None or empty
    return ""

# Dictionary untuk hari dan bulan dalam bahasa Indonesia
hari = {
    'Monday': 'Senin', 'Tuesday': 'Selasa', 'Wednesday': 'Rabu', 'Thursday': 'Kamis',
    'Friday': 'Jumat', 'Saturday': 'Sabtu', 'Sunday': 'Minggu'
}
bulan = {
    'January': 'Januari', 'February': 'Februari', 'March': 'Maret', 'April': 'April',
    'May': 'Mei', 'June': 'Juni', 'July': 'Juli', 'August': 'Agustus',
    'September': 'September', 'October': 'Oktober', 'November': 'November', 'December': 'Desember'
}

# Fungsi untuk mengubah dokumen Word menjadi file yang dapat diunduh
def create_download_link(docx_file, filename):
    # Simpan dokumen ke BytesIO
    doc_io = io.BytesIO()
    docx_file.save(doc_io)
    doc_io.seek(0)
    
    # Encode untuk download link
    b64 = base64.b64encode(doc_io.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download {filename}</a>'

# Fungsi untuk mengubah dokumen Word menjadi PDF menggunakan LibreOffice
def convert_docx_to_pdf(docx_file):
    try:
        st.info("Mengkonversi dokumen ke PDF menggunakan LibreOffice...")
        
        # Simpan dokumen ke file sementara
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            docx_file.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        # Nama file PDF sementara
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        # Konversi dengan LibreOffice
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', 
            '--outdir', os.path.dirname(temp_pdf_path), temp_docx_path
        ], check=True, timeout=30)
        
        # Baca file PDF yang dihasilkan
        with open(temp_pdf_path, 'rb') as pdf_file:
            pdf_data = pdf_file.read()
        
        # Hapus file sementara
        os.unlink(temp_docx_path)
        os.unlink(temp_pdf_path)
        
        return pdf_data
    except Exception as e:
        st.error(f"Konversi PDF gagal: {str(e)}")
        
        # Pastikan file sementara dihapus
        if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
            try:
                os.unlink(temp_docx_path)
            except:
                pass
        if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
            try:
                os.unlink(temp_pdf_path)
            except:
                pass
                
        return None

# Fungsi untuk membuat link download PDF
def create_pdf_download_link(pdf_data, filename):
    b64 = base64.b64encode(pdf_data).decode()
    return f'<a href="data:application/pdf;base64,{b64}" download="{filename}">Download {filename}</a>'

# Main function untuk streamlit
def generate_report():
    # Ambil tanggal saat ini
    sekarang = datetime.now()
    nama_hari = hari[sekarang.strftime("%A")]
    nama_bulan = bulan[sekarang.strftime("%B")]
    tanggal = sekarang.strftime("%d")
    tahun = sekarang.strftime("%Y")
    waktu_laporan = f"{nama_hari}, {tanggal} {nama_bulan} {tahun}"
    
    # Nama file untuk download
    file_name = f"{tanggal} {nama_bulan} {tahun}_Daily Report Wildan Dzaky Ramadhani.docx"
    
    # Tanggal filter (untuk default gunakan tanggal saat ini)
    filter_date = st.date_input("Pilih tanggal untuk laporan:", sekarang)
    filter_tanggal = filter_date.strftime("%Y-%m-%d")
    
    if st.button("Generate Report"):
        with st.spinner('Generating report...'):
            # URL ekspor Google Sheets untuk file Excel
            file_id = "1wJlAUerJDxpaBRxMOLxOmcSzG5LdwUz4K8HJ3uc1v0s"
            export_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx"

            # Unduh file Excel langsung dari Google Sheets
            try:
                response = requests.get(export_url)
                if response.status_code == 200:
                    excel_data = io.BytesIO(response.content)
                    df = pd.read_excel(excel_data, sheet_name="New Format")
                else:
                    st.error(f"Gagal mengunduh file: status code {response.status_code}")
                    return
            except Exception as e:
                st.error(f"Gagal mengunduh file: {str(e)}")
                return

            # Filter data berdasarkan tanggal
            df_filtered = df[df['Tanggal'] == filter_tanggal].fillna('')  # Ganti NaN dengan string kosong

            if df_filtered.empty:
                st.warning(f"Tidak ada data untuk tanggal {filter_tanggal}")
                return

            # Konversi data filtered ke format yang sesuai untuk tabel
            data_baru = [
                (
                    idx + 1,  # Override nomor urut agar dimulai dari 1
                    row['Pekerjaan'] if not pd.isna(row['Pekerjaan']) else "",
                    format_tanggal(str(row['Batas Waktu']), bulan),
                    row['Status'] if not pd.isna(row['Status']) else "",
                    format_tanggal(str(row['Diselesaikan Pada']), bulan)
                )
                for idx, (_, row) in enumerate(df_filtered.iterrows())
            ]

            # Ambil semua keterangan dari data yang difilter dan gabungkan
            keterangan_list = [str(row['Keterangan']).strip() for _, row in df_filtered.iterrows() if str(row['Keterangan']).strip()]
            keterangan = "\n".join(keterangan_list) if keterangan_list else "Tidak ada keterangan untuk hari ini."

            # Cek apakah template dokumen ada
            template_path = "Weekly Daily Report Wildan Dzaky Ramadhani.docx"
            if not os.path.exists(template_path):
                st.error(f"Template file '{template_path}' tidak ditemukan.")
                return
                
            # Baca dokumen yang sudah ada
            doc = Document(template_path)

            # Update Waktu Laporan saja
            for para in doc.paragraphs:
                if "Waktu Laporan" in para.text:
                    para.text = f"Waktu Laporan\t: {waktu_laporan}"

            # Hapus paragraf Catatan jika ada
            paragraphs_to_remove = []
            for i, para in enumerate(doc.paragraphs):
                if "Catatan" in para.text:
                    paragraphs_to_remove.append(para)
            
            for para in paragraphs_to_remove:
                doc.element.body.remove(para._element)
                
            # PENDEKATAN BARU: Kita tidak menghapus tabel lama, tapi menggantinya dengan data baru
            if not doc.tables:
                st.error("Tidak dapat menemukan tabel di template dokumen.")
                return
                
            table = doc.tables[0]  # Gunakan tabel yang ada
            
            # Pertahankan baris header, hapus semua baris data (baris ke-2 dan seterusnya)
            # Ini mempertahankan semua format dan style
            while len(table.rows) > 1:
                tr = table.rows[1]._tr
                table._tbl.remove(tr)
                
                
            # Tambahkan data baru ke tabel yang ada
            for no, pekerjaan, batas_waktu, status, selesai_pada in data_baru:
                row_cells = table.add_row().cells
                row_cells[0].text = str(no)
                row_cells[1].text = pekerjaan
                row_cells[2].text = batas_waktu
                row_cells[3].text = status
                row_cells[4].text = selesai_pada
                
                # Tambahkan garis hitam saja, tidak perlu menyalin format warna
                # karena kita menggunakan tabel asli yang sudah memiliki format
                for i in range(len(row_cells)):
                    add_black_borders(row_cells[i])

            # Tambahkan Keterangan setelah tabel
            doc.add_paragraph(f"{keterangan}")

            # Buat download link untuk DOCX
            download_html = create_download_link(doc, file_name)
            st.markdown(download_html, unsafe_allow_html=True)
            
            # Konversi ke PDF jika diminta
            if st.checkbox("Buat juga versi PDF"):
                with st.spinner("Mengkonversi ke PDF..."):
                    pdf_data = convert_docx_to_pdf(doc)
                    if pdf_data:
                        pdf_file_name = file_name.replace('.docx', '.pdf')
                        pdf_download_html = create_pdf_download_link(pdf_data, pdf_file_name)
                        st.markdown(pdf_download_html, unsafe_allow_html=True)
                        st.success(f'Report berhasil dibuat dalam format DOCX dan PDF! Klik link di atas untuk mengunduh.')
                    else:
                        st.warning("Konversi PDF gagal. Hanya file DOCX yang tersedia.")
                        st.success(f'Report berhasil dibuat! Klik link di atas untuk mengunduh file DOCX.')
            else:
                st.success(f'Report berhasil dibuat! Klik link di atas untuk mengunduh.')

if __name__ == "__main__":
    generate_report()

# -------------------------------------------------------------------------------------------
# PANDUAN INSTALASI DAN MENJALANKAN APLIKASI DI LINUX UBUNTU
# -------------------------------------------------------------------------------------------
# Jalankan perintah berikut di terminal Ubuntu untuk instalasi:
#
# 1. Update sistem dan instal dependencies
# sudo apt update
# sudo apt install -y python3 python3-pip python3-venv libreoffice
#
# 2. Buat direktori untuk aplikasi
# mkdir -p ~/daily_report_app
#
# 3. Pindahkan file ini ke direktori aplikasi
# cp streamlit.py ~/daily_report_app/
# cp "Weekly Daily Report Wildan Dzaky Ramadhani.docx" ~/daily_report_app/
#
# 4. Buat dan aktifkan virtual environment
# cd ~/daily_report_app
# python3 -m venv venv
# source venv/bin/activate
#
# 5. Instal library yang diperlukan
# pip install streamlit python-docx pandas requests
#
# 6. Jalankan aplikasi agar bisa diakses melalui IP:
# streamlit run streamlit.py --server.address=0.0.0.0 --server.port=8501
#
# 7. Akses aplikasi dari browser:
# http://<IP-SERVER>:8501
#
# Catatan: Ganti <IP-SERVER> dengan IP komputer/server Ubuntu Anda.
# Jika menggunakan UFW firewall, buka port dengan: sudo ufw allow 8501/tcp
# 
# Untuk menjalankan aplikasi sebagai layanan yang tetap berjalan:
# 1. Buat file service systemd:
#    sudo nano /etc/systemd/system/daily-report.service
#
# 2. Isi dengan konfigurasi berikut:
#    [Unit]
#    Description=Daily Report Streamlit App
#    After=network.target
#    
#    [Service]
#    User=<username>
#    WorkingDirectory=/home/<username>/daily_report_app
#    ExecStart=/home/<username>/daily_report_app/venv/bin/streamlit run streamlit.py --server.address=0.0.0.0 --server.port=8501
#    Restart=always
#    RestartSec=5
#    
#    [Install]
#    WantedBy=multi-user.target
#
# 3. Ganti <username> dengan nama pengguna Anda
#
# 4. Aktifkan dan jalankan service:
#    sudo systemctl daemon-reload
#    sudo systemctl enable daily-report
#    sudo systemctl start daily-report
#
# 5. Cek status service:
#    sudo systemctl status daily-report
from flask import Flask, render_template, request, send_file, Response, jsonify
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import requests
import io
import base64
import os
import sys
import logging
from werkzeug.utils import secure_filename
import tempfile
import subprocess
import traceback

# Configure logging - Console only to avoid feedback loops
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
        # Removed FileHandler to avoid watchdog feedback loops
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Create templates directory if it doesn't exist
os.makedirs('templates', exist_ok=True)

# Create a simple index.html file if it doesn't exist
if not os.path.exists('templates/index.html'):
    with open('templates/index.html', 'w') as f:
        f.write('''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Daily Report Generator</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .container { max-width: 600px; margin: 0 auto; }
                .form-group { margin-bottom: 15px; }
                label { display: block; margin-bottom: 5px; }
                input, select { padding: 8px; width: 100%; }
                button { padding: 10px 15px; background: #4CAF50; color: white; border: none; cursor: pointer; }
                .error { color: red; margin-top: 20px; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Daily Report Generator</h1>
                {% if error %}
                <div class="error">
                    <p>Error: {{ error }}</p>
                </div>
                {% endif %}
                <form method="POST" enctype="multipart/form-data">
                    <div class="form-group">
                        <label for="filter_date">Select Date:</label>
                        <input type="date" id="filter_date" name="filter_date" value="{{ date }}" required>
                    </div>
                    <div class="form-group">
                        <label for="format_type">Output Format:</label>
                        <select id="format_type" name="format_type">
                            <option value="docx">DOCX</option>
                            <option value="pdf">PDF</option>
                        </select>
                    </div>
                    <button type="submit">Generate Report</button>
                </form>
            </div>
        </body>
        </html>
        ''')

# Fungsi untuk menyalin pemformatan warna sel
def copy_cell_formatting(src_cell, dest_cell):
    try:
        tcPr = src_cell._tc.get_or_add_tcPr()
        shd = tcPr.first_child_found_in("w:shd")
        if shd is not None:
            new_tcPr = dest_cell._tc.get_or_add_tcPr()
            new_shd = OxmlElement('w:shd')
            new_shd.set(qn('w:fill'), shd.get(qn('w:fill')))
            new_shd.set(qn('w:val'), shd.get(qn('w:val')))
            new_tcPr.append(new_shd)
    except Exception as e:
        logger.warning(f"Error copying cell formatting: {str(e)}")

# Fungsi untuk menambahkan garis hitam pada sel
def add_black_borders(cell):
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    except Exception as e:
        logger.warning(f"Error adding borders: {str(e)}")

# Fungsi untuk mengubah format tanggal
def format_tanggal(tanggal_str, bulan_dict):
    if isinstance(tanggal_str, str) and tanggal_str.strip():
        try:
            # Check for "NaT" string or empty string
            if tanggal_str == "NaT" or tanggal_str == "":
                return ""
                
            dt = datetime.strptime(tanggal_str.split()[0], "%Y-%m-%d")
            hari = dt.strftime("%d")
            bulan = bulan_dict[dt.strftime("%B")]
            tahun = dt.strftime("%Y")
            return f"{hari} {bulan} {tahun}"
        except ValueError as e:
            logger.warning(f"Error formatting date {tanggal_str}: {str(e)}")
            return ""
    return ""

# Dictionary untuk hari dan bulan dalam bahasa Indonesia
hari = {
    'Monday': 'Senin', 'Tuesday': 'Selasa', 'Wednesday': 'Rabu', 'Thursday': 'Kamis',
    'Friday': 'Jumat', 'Saturday': 'Sabtu', 'Sunday': 'Minggu'
}
bulan = {
    'January': 'Januari', 'February': 'Februari', 'March': 'Maret', 'April': 'April',
    'May': 'Mei', 'June': 'Juni', 'July': 'Juli', 'August': 'Agustus',
    'September': 'September', 'October': 'Oktober', 'November': 'November', 'December': 'Desember'
}

def generate_docx_report(filter_date):
    logger.info(f"Generating report for date: {filter_date}")
    
    # Ambil tanggal dari parameter
    sekarang = filter_date if filter_date else datetime.now()
    nama_hari = hari[sekarang.strftime("%A")]
    nama_bulan = bulan[sekarang.strftime("%B")]
    tanggal = sekarang.strftime("%d")
    tahun = sekarang.strftime("%Y")
    waktu_laporan = f"{nama_hari}, {tanggal} {nama_bulan} {tahun}"
    
    # Format untuk filter
    filter_tanggal = sekarang.strftime("%Y-%m-%d")
    
    # Nama file untuk download
    file_name = f"{tanggal} {nama_bulan} {tahun}_Daily Report Wildan Dzaky Ramadhani.docx"
    
    # URL ekspor Google Sheets untuk file Excel
    file_id = "1wJlAUerJDxpaBRxMOLxOmcSzG5LdwUz4K8HJ3uc1v0s"
    export_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx"

    # Unduh file Excel langsung dari Google Sheets
    try:
        logger.info(f"Downloading spreadsheet from Google Sheets: {file_id}")
        response = requests.get(export_url)
        if response.status_code != 200:
            error_msg = f"Gagal mengunduh file: status code {response.status_code}"
            logger.error(error_msg)
            return {"error": error_msg}, None
        excel_data = io.BytesIO(response.content)
        logger.info("Spreadsheet downloaded successfully")
        
        logger.info("Reading Excel data")
        df = pd.read_excel(excel_data, sheet_name="New Format")
        logger.info(f"Excel data read successfully, shape: {df.shape}")
    except Exception as e:
        error_msg = f"Gagal mengunduh file: {str(e)}"
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        return {"error": error_msg}, None

    # Filter data berdasarkan tanggal
    logger.info(f"Filtering data for date: {filter_tanggal}")
    df_filtered = df[df['Tanggal'] == filter_tanggal].fillna('')
    logger.info(f"Filtered data shape: {df_filtered.shape}")

    if df_filtered.empty:
        error_msg = f"Tidak ada data untuk tanggal {filter_tanggal}"
        logger.warning(error_msg)
        return {"error": error_msg}, None

    # Konversi data filtered ke format yang sesuai untuk tabel
    data_baru = [
        (
            idx + 1,
            row['Pekerjaan'] if not pd.isna(row['Pekerjaan']) else "",
            format_tanggal(str(row['Batas Waktu']), bulan),
            row['Status'] if not pd.isna(row['Status']) else "",
            format_tanggal(str(row['Diselesaikan Pada']), bulan)
        )
        for idx, (_, row) in enumerate(df_filtered.iterrows())
    ]

    # Ambil semua keterangan dari data yang difilter dan gabungkan
    keterangan_list = [str(row['Keterangan']).strip() for _, row in df_filtered.iterrows() if str(row['Keterangan']).strip()]
    keterangan = "\n".join(keterangan_list) if keterangan_list else "Tidak ada keterangan untuk hari ini."

    # Cek apakah template dokumen ada
    template_path = "Weekly Daily Report Wildan Dzaky Ramadhani.docx"
    
    # If template doesn't exist, create a simple one
    if not os.path.exists(template_path):
        logger.warning(f"Template file '{template_path}' not found. Creating a simple template.")
        doc = Document()
        doc.add_heading('Daily Report', 0)
        doc.add_paragraph(f"Waktu Laporan\t: {waktu_laporan}")
        
        # Add table with headers
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ["No", "Pekerjaan", "Batas Waktu", "Status", "Diselesaikan Pada"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
    else:
        try:
            logger.info(f"Opening template file: {template_path}")
            doc = Document(template_path)
            logger.info("Template file opened successfully")
        except Exception as e:
            error_msg = f"Gagal membuka template dokumen: {str(e)}"
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            return {"error": error_msg}, None

    # Update Waktu Laporan saja
    for para in doc.paragraphs:
        if "Waktu Laporan" in para.text:
            para.text = f"Waktu Laporan\t: {waktu_laporan}"
            break
    else:
        # If not found, add it
        doc.add_paragraph(f"Waktu Laporan\t: {waktu_laporan}")

    # Hapus paragraf Catatan jika ada
    paragraphs_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        if "Catatan" in para.text:
            paragraphs_to_remove.append(para)
    
    for para in paragraphs_to_remove:
        try:
            doc.element.body.remove(para._element)
        except Exception as e:
            logger.warning(f"Error removing paragraph: {str(e)}")
            
    # Check if table exists, if not create one
    if not doc.tables:
        logger.warning("No table found in template, creating one")
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ["No", "Pekerjaan", "Batas Waktu", "Status", "Diselesaikan Pada"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
    else:
        table = doc.tables[0]
            
    # Pertahankan baris header, hapus semua baris data (baris ke-2 dan seterusnya)
    while len(table.rows) > 1:
        try:
            tr = table.rows[1]._tr
            table._tbl.remove(tr)
        except Exception as e:
            logger.warning(f"Error removing table row: {str(e)}")
            break
            
    # Tambahkan data baru ke tabel yang ada
    logger.info(f"Adding {len(data_baru)} rows to table")
    for no, pekerjaan, batas_waktu, status, selesai_pada in data_baru:
        try:
            row_cells = table.add_row().cells
            row_cells[0].text = str(no)
            row_cells[1].text = pekerjaan
            row_cells[2].text = batas_waktu
            row_cells[3].text = status
            row_cells[4].text = selesai_pada
            
            # Tambahkan garis hitam
            for i in range(len(row_cells)):
                add_black_borders(row_cells[i])
        except Exception as e:
            logger.error(f"Error adding row to table: {str(e)}")
            logger.error(traceback.format_exc())

    # Tambahkan Keterangan setelah tabel
    doc.add_paragraph(f"{keterangan}")

    # Simpan ke BytesIO
    logger.info("Saving document to memory")
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    logger.info("Document saved successfully")
    
    return {"success": True, "file_name": file_name}, docx_io

def convert_to_pdf(docx_content, base_filename):
    """
    Convert DOCX content to PDF with improved error handling
    """
    logger.info("Starting PDF conversion")
    
    # Create temporary files with unique names to avoid conflicts
    temp_dir = tempfile.mkdtemp()
    temp_docx_path = os.path.join(temp_dir, f"{base_filename}.docx")
    
    with open(temp_docx_path, 'wb') as f:
        f.write(docx_content)
    
    logger.info(f"Created temporary DOCX file: {temp_docx_path}")
    
    # Define output PDF path
    temp_pdf_path = os.path.join(temp_dir, f"{base_filename}.pdf")
    
    try:
        # Set environment variables for LibreOffice
        env = os.environ.copy()
        if 'HOME' not in env:
            env['HOME'] = '/root'
        
        # Use absolute path to LibreOffice
        libreoffice_path = '/usr/bin/libreoffice'
        
        logger.info(f"Running LibreOffice conversion: {temp_docx_path} -> {temp_pdf_path}")
        
        # Execute the command with detailed output capture
        process = subprocess.run(
            [
                libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                temp_docx_path
            ],
            env=env,
            capture_output=True,
            text=True,
            timeout=180  # 3 minutes timeout
        )
        
        logger.info(f"Command output: {process.stdout}")
        if process.stderr:
            logger.info(f"Command stderr: {process.stderr}")
        
        # Verify the PDF was created
        if not os.path.exists(temp_pdf_path):
            raise Exception("PDF file was not created after conversion")
        
        # Read the PDF content
        with open(temp_pdf_path, 'rb') as f:
            pdf_content = io.BytesIO(f.read())
            pdf_content.seek(0)
        
        return pdf_content
        
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out")
        raise Exception("PDF conversion timed out. Try downloading as DOCX instead.")
        
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed with exit code {e.returncode}")
        logger.error(f"STDOUT: {e.stdout}")
        logger.error(f"STDERR: {e.stderr}")
        raise Exception(f"PDF conversion failed with exit code {e.returncode}. Try downloading as DOCX instead.")
        
    except Exception as e:
        logger.error(f"PDF conversion failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise
        
    finally:
        # Clean up temporary files
        try:
            import shutil
            shutil.rmtree(temp_dir)
            logger.info(f"Cleaned up temporary directory: {temp_dir}")
        except Exception as e:
            logger.warning(f"Error cleaning up temporary files: {str(e)}")

@app.route('/status')
def status():
    """Simple status endpoint to check if app is running"""
    libreoffice_info = check_libreoffice_availability()
    return jsonify({
        "status": "running",
        "template_exists": os.path.exists("Weekly Daily Report Wildan Dzaky Ramadhani.docx"),
        "templates_dir_exists": os.path.exists("templates"),
        "index_html_exists": os.path.exists("templates/index.html"),
        "libreoffice": libreoffice_info
    })

def check_libreoffice_availability():
    """Check if LibreOffice is available and working"""
    try:
        result = subprocess.run(['/usr/bin/libreoffice', '--version'], 
                               capture_output=True, 
                               text=True, 
                               timeout=10)
        if result.returncode == 0:
            return {"available": True, "path": "/usr/bin/libreoffice", "version": result.stdout.strip()}
    except Exception as e:
        return {"available": False, "error": str(e)}
    
    return {"available": False}

@app.route('/debug')
def debug():
    """Endpoint to get debug information"""
    import platform
    import sys
    
    debug_info = {
        "python_version": sys.version,
        "platform": platform.platform(),
        "working_directory": os.getcwd(),
        "files_in_directory": os.listdir(),
        "libreoffice_status": check_libreoffice_availability(),
        "environment_variables": {k: v for k, v in os.environ.items() 
                                 if not k.lower().startswith(('pass', 'secret', 'key', 'token'))}
    }
    
    return jsonify(debug_info)

@app.route('/', methods=['GET', 'POST'])
def index():
    try:
        if request.method == 'POST':
            logger.info("Received POST request")
            filter_date_str = request.form.get('filter_date')
            format_type = request.form.get('format_type', 'docx')
            
            logger.info(f"Request parameters: date={filter_date_str}, format={format_type}")
            
            # Parse date string to datetime object
            try:
                filter_date = datetime.strptime(filter_date_str, "%Y-%m-%d")
            except Exception as e:
                logger.warning(f"Invalid date format: {str(e)}. Using current date.")
                filter_date = datetime.now()
            
            # Generate the DOCX report first
            result, docx_io = generate_docx_report(filter_date)
            
            if 'error' in result:
                return render_template('index.html', error=result['error'], date=datetime.now().strftime("%Y-%m-%d"))
            
            # Generate file name
            nama_hari = hari[filter_date.strftime("%A")]
            nama_bulan = bulan[filter_date.strftime("%B")]
            tanggal = filter_date.strftime("%d")
            tahun = filter_date.strftime("%Y")
            base_filename = f"{tanggal} {nama_bulan} {tahun}_Daily Report Wildan Dzaky Ramadhani"
            
            logger.info(f"Preparing to send file: {base_filename}.{format_type}")
            
            # Return according to format requested
            if format_type == 'docx':
                # Send the DOCX file directly
                return send_file(
                    docx_io, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f"{base_filename}.docx"
                )
            elif format_type == 'pdf':
                try:
                    # Use the improved PDF conversion function
                    pdf_io = convert_to_pdf(docx_io.getvalue(), base_filename)
                    
                    # Send the PDF file
                    return send_file(
                        pdf_io,
                        mimetype='application/pdf',
                        as_attachment=True,
                        download_name=f"{base_filename}.pdf"
                    )
                except Exception as e:
                    logger.error(f"Error in PDF conversion: {str(e)}")
                    return render_template('index.html', 
                                          error=f"Error converting to PDF. Please try DOCX format instead. Technical details: {str(e)}",
                                          date=datetime.now().strftime("%Y-%m-%d"))
        
        # GET request or initial page load
        logger.info("Handling GET request")
        return render_template('index.html', date=datetime.now().strftime("%Y-%m-%d"))
    
    except Exception as e:
        logger.error(f"Unhandled exception: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"error": str(e), "traceback": traceback.format_exc()})

if __name__ == '__main__':
    print("Starting Flask application...")
    print(f"Current directory: {os.getcwd()}")
    print(f"Python version: {sys.version}")
    print(f"Template exists: {os.path.exists('Weekly Daily Report Wildan Dzaky Ramadhani.docx')}")
    print(f"Templates directory exists: {os.path.exists('templates')}")
    
    # Check LibreOffice availability
    libreoffice_status = check_libreoffice_availability()
    print(f"LibreOffice status: {libreoffice_status}")
    
    # Create a test docx if template doesn't exist (for demo purposes)
    if not os.path.exists("Weekly Daily Report Wildan Dzaky Ramadhani.docx"):
        try:
            doc = Document()
            doc.add_heading('Daily Report Template', 0)
            doc.add_paragraph("Waktu Laporan\t: Template")
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            for i, header in enumerate(["No", "Pekerjaan", "Batas Waktu", "Status", "Diselesaikan Pada"]):
                header_cells[i].text = header
            doc.save("Weekly Daily Report Wildan Dzaky Ramadhani.docx")
            print("Created a sample template docx file")
        except Exception as e:
            print(f"Failed to create sample template: {e}")
    
    app.run(host='0.0.0.0', port=5000, debug=True)